import telebot
import datetime
import os
import io
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import Flask, request
import threading
import time

# TOKEN
TOKEN = "8716483413:AAEqs14n3l6hRrXC6hwAhkggTqI772ub-iY"

bot = telebot.TeleBot(TOKEN)
app = Flask(__name__)

agents = {}
visits = []  # Doktor tashriflari
pharmacy_visits = []  # Apteka tashriflari

# Webhook uchun route
@app.route('/webhook', methods=['POST'])
def webhook():
    json_str = request.get_data().decode('UTF-8')
    update = telebot.types.Update.de_json(json_str)
    bot.process_new_updates([update])
    return 'OK', 200

# Bot polling o'rniga webhook bilan ishlaydi
def setup_webhook():
    time.sleep(2)  # Render to'liq ishga tushishi uchun
    try:
        # Render URL'ingizni yozing (keyin aniqlaymiz)
        render_url = f"https://{os.environ.get('RENDER_EXTERNAL_URL', 'kpi-med-bot.onrender.com')}"
        bot.remove_webhook()
        bot.set_webhook(url=f"{render_url}/webhook")
        print(f"✅ Webhook sozlandi: {render_url}/webhook")
    except Exception as e:
        print(f"❌ Webhook xatosi: {e}")

# ============== OYLIK HISOBOT FUNKSIYALARI ==============

def generate_excel_report(month=None, year=None):
    """Excel formatda oylik hisobot yaratish"""
    today = datetime.date.today()
    if month is None:
        month = today.month
    if year is None:
        year = today.year
    
    # Oy boshi va oxiri
    month_start = datetime.date(year, month, 1)
    if month == 12:
        month_end = datetime.date(year + 1, 1, 1) - datetime.timedelta(days=1)
    else:
        month_end = datetime.date(year, month + 1, 1) - datetime.timedelta(days=1)
    
    # Excel fayl yaratish
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"{month}-{year} Hisobot"
    
    # Sarlavha
    ws.merge_cells('A1:G1')
    title_cell = ws['A1']
    title_cell.value = f"OYLIK HISOBOT ({month}-OY {year})"
    title_cell.font = Font(size=16, bold=True)
    title_cell.alignment = Alignment(horizontal='center')
    
    # Sarlavha qatorlari
    headers = [
        ['Agent', 'Doktor', 'Apteka', 'Doktor KPI', 'Apteka KPI', 'Umumiy KPI', 'Reyting'],
        ['', '(ta)', '(ta)', '(%)', '(%)', '(%)', '']
    ]
    
    for row_idx, header_row in enumerate(headers, start=3):
        for col_idx, header in enumerate(header_row, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal='center')
    
    # Agentlar bo'yicha statistika yig'ish
    agents_stats = {}
    
    # Doktorlar
    for v in visits:
        if month_start <= v["date"] <= month_end:
            user_id = v["user"]
            if user_id not in agents_stats:
                agents_stats[user_id] = {
                    "name": v.get("user_name", f"User {user_id}"),
                    "doctor_count": 0,
                    "pharmacy_count": 0
                }
            agents_stats[user_id]["doctor_count"] += 1
    
    # Aptekalar
    for p in pharmacy_visits:
        if month_start <= p["date"] <= month_end:
            user_id = p["user"]
            if user_id in agents_stats:
                agents_stats[user_id]["pharmacy_count"] += 1
    
    # Ma'lumotlarni yozish
    row_num = 5
    for agent_id, data in agents_stats.items():
        # Kunlar soni
        days_in_month = month_end.day
        
        # KPI hisoblash
        doctor_kpi = min(100, (data["doctor_count"] / (days_in_month * 12)) * 100)
        apteka_kpi = min(100, (data["pharmacy_count"] / (days_in_month * 5)) * 100)
        total_kpi = (doctor_kpi * 0.7) + (apteka_kpi * 0.3)
        
        ws.cell(row=row_num, column=1, value=data["name"])
        ws.cell(row=row_num, column=2, value=data["doctor_count"])
        ws.cell(row=row_num, column=3, value=data["pharmacy_count"])
        ws.cell(row=row_num, column=4, value=round(doctor_kpi, 1))
        ws.cell(row=row_num, column=5, value=round(apteka_kpi, 1))
        ws.cell(row=row_num, column=6, value=round(total_kpi, 1))
        
        # KPI bo'yicha rang berish
        kpi_cell = ws.cell(row=row_num, column=6)
        if total_kpi >= 90:
            kpi_cell.fill = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")  # Yashil
        elif total_kpi >= 70:
            kpi_cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")  # Sariq
        else:
            kpi_cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")  # Qizil
            kpi_cell.font = Font(color="FFFFFF")
        
        row_num += 1
    
    # Reyting bo'yicha tartiblash
    data_rows = []
    for row in range(5, row_num):
        data_rows.append([ws.cell(row=row, col=1).value,  # Agent
                         ws.cell(row=row, col=6).value])  # KPI
    
    # Reyting yozish
    data_rows.sort(key=lambda x: x[1], reverse=True)
    for idx, (agent, kpi) in enumerate(data_rows, 1):
        for row in range(5, row_num):
            if ws.cell(row=row, column=1).value == agent:
                ws.cell(row=row, column=7, value=idx)
                break
    
    # Formatlash
    for col in range(1, 8):
        col_letter = openpyxl.utils.get_column_letter(col)
        ws.column_dimensions[col_letter].width = 15
    
    # Excel faylni memory ga saqlash
    excel_file = io.BytesIO()
    wb.save(excel_file)
    excel_file.seek(0)
    
    return excel_file, f"oylik_hisobot_{month}_{year}.xlsx"


def generate_word_report(month=None, year=None):
    """Word formatda oylik hisobot yaratish"""
    today = datetime.date.today()
    if month is None:
        month = today.month
    if year is None:
        year = today.year
    
    month_start = datetime.date(year, month, 1)
    if month == 12:
        month_end = datetime.date(year + 1, 1, 1) - datetime.timedelta(days=1)
    else:
        month_end = datetime.date(year, month + 1, 1) - datetime.timedelta(days=1)
    
    # Word document yaratish
    doc = Document()
    
    # Sarlavha
    title = doc.add_heading(f'OYLIK HISOBOT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(f'{month}-OY {year} yil', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Agentlar statistikasi
    agents_stats = {}
    for v in visits:
        if month_start <= v["date"] <= month_end:
            user_id = v["user"]
            if user_id not in agents_stats:
                agents_stats[user_id] = {
                    "name": v.get("user_name", f"User {user_id}"),
                    "doctor_count": 0,
                    "pharmacy_count": 0
                }
            agents_stats[user_id]["doctor_count"] += 1
    
    for p in pharmacy_visits:
        if month_start <= p["date"] <= month_end:
            user_id = p["user"]
            if user_id in agents_stats:
                agents_stats[user_id]["pharmacy_count"] += 1
    
    days_in_month = month_end.day
    
    # Jadval yaratish
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Sarlavha qatori
    header_cells = table.rows[0].cells
    headers = ['Agent', 'Doktor', 'Apteka', 'Doktor KPI', 'Apteka KPI', 'Umumiy KPI', 'Reyting']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    
    # Ma'lumotlar
    agents_list = []
    for agent_id, data in agents_stats.items():
        doctor_kpi = min(100, (data["doctor_count"] / (days_in_month * 12)) * 100)
        apteka_kpi = min(100, (data["pharmacy_count"] / (days_in_month * 5)) * 100)
        total_kpi = (doctor_kpi * 0.7) + (apteka_kpi * 0.3)
        
        agents_list.append({
            "name": data["name"],
            "doctor": data["doctor_count"],
            "pharmacy": data["pharmacy_count"],
            "doctor_kpi": round(doctor_kpi, 1),
            "apteka_kpi": round(apteka_kpi, 1),
            "total_kpi": round(total_kpi, 1)
        })
    
    # Reyting bo'yicha tartiblash
    agents_list.sort(key=lambda x: x["total_kpi"], reverse=True)
    
    for idx, agent in enumerate(agents_list, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = agent["name"]
        row_cells[1].text = str(agent["doctor"])
        row_cells[2].text = str(agent["pharmacy"])
        row_cells[3].text = f"{agent['doctor_kpi']}%"
        row_cells[4].text = f"{agent['apteka_kpi']}%"
        row_cells[5].text = f"{agent['total_kpi']}%"
        row_cells[6].text = str(idx)
    
    # Umumiy statistika
    doc.add_paragraph()
    doc.add_heading('Umumiy statistika', level=2)
    
    total_doctors = sum(a["doctor"] for a in agents_list)
    total_pharmacies = sum(a["pharmacy"] for a in agents_list)
    avg_kpi = sum(a["total_kpi"] for a in agents_list) / len(agents_list) if agents_list else 0
    
    stats_para = doc.add_paragraph()
    stats_para.add_run(f'Jami doktor tashriflari: ').bold = True
    stats_para.add_run(f'{total_doctors}\n')
    stats_para.add_run(f'Jami apteka tashriflari: ').bold = True
    stats_para.add_run(f'{total_pharmacies}\n')
    stats_para.add_run(f'O\'rtacha KPI: ').bold = True
    stats_para.add_run(f'{avg_kpi:.1f}%\n')
    
    # Eng yaxshi agent
    if agents_list:
        doc.add_paragraph()
        best = agents_list[0]
        best_para = doc.add_paragraph()
        best_para.add_run('🏆 ENG YAXSHI AGENT: ').bold = True
        best_para.add_run(f"{best['name']} ({best['total_kpi']}% KPI)")
    
    # Word faylni memory ga saqlash
    word_file = io.BytesIO()
    doc.save(word_file)
    word_file.seek(0)
    
    return word_file, f"oylik_hisobot_{month}_{year}.docx"


# ============== BOT BUYRUQLARI ==============

@bot.message_handler(commands=['start'])
def start(message):
    welcome_text = """
🤖 **KPI BOT**

Assalomu alaykum! KPI bot orqali kunlik faoliyatingizni kuzatishingiz mumkin.

📋 **MAVJUD BUYRUQLAR:**

/help - Barcha buyruqlar ro'yxati
/doctor - Doktor tashrifini qo'shish
/apteka - Apteka tashrifini qo'shish
/report - Bugungi hisobot
/monthly_excel - Oylik hisobot (Excel formatda)
/monthly_word - Oylik hisobot (Word formatda)
/plan - Ertangi reja tuzish
/stats - Oylik statistika

✅ Foydalanish uchun yuqoridagi buyruqlardan birini tanlang.
    """
    bot.send_message(message.chat.id, welcome_text, parse_mode='Markdown')


@bot.message_handler(commands=['help'])
def help_command(message):
    help_text = """
📋 **BATAFSIL BUYRUQLAR:**

/start - Botni ishga tushirish
/help - Bu yordam oynasi

**🏥 ASOSIY BUYRUQLAR:**
/doctor - Doktor tashrifini qo'shish
/apteka - Apteka tashrifini qo'shish
/plan - Ertangi kun rejasini tuzish
/report - Bugungi kun hisoboti
/stats - Oylik statistika

**📊 HISOBOTLAR:**
/monthly_excel - Excel formatda oylik hisobot
/monthly_word - Word formatda oylik hisobot

**📊 MANAGER UCHUN:**
/all_report - Barcha agentlar hisoboti

✅ Doktor normasi: 12 ta/kun
✅ Apteka normasi: 5 ta/kun
🔹 Har bir doktorga kuniga 1 marta kirish mumkin
    """
    bot.send_message(message.chat.id, help_text, parse_mode='Markdown')


@bot.message_handler(commands=['monthly_excel'])
def monthly_excel(message):
    if message.from_user.id != 6995658625:
        bot.send_message(message.chat.id, "⛔ Bu buyruq faqat manager uchun!")
        return
    
    bot.send_message(message.chat.id, "⏳ Excel hisobot tayyorlanmoqda...")
    
    try:
        today = datetime.date.today()
        excel_file, filename = generate_excel_report(today.month, today.year)
        
        bot.send_document(
            message.chat.id,
            document=excel_file,
            visible_file_name=filename,
            caption=f"📊 Oylik hisobot ({today.month}-oy {today.year})"
        )
    except Exception as e:
        bot.send_message(message.chat.id, f"❌ Xatolik: {str(e)}")


@bot.message_handler(commands=['monthly_word'])
def monthly_word(message):
    if message.from_user.id != 6995658625:
        bot.send_message(message.chat.id, "⛔ Bu buyruq faqat manager uchun!")
        return
    
    bot.send_message(message.chat.id, "⏳ Word hisobot tayyorlanmoqda...")
    
    try:
        today = datetime.date.today()
        word_file, filename = generate_word_report(today.month, today.year)
        
        bot.send_document(
            message.chat.id,
            document=word_file,
            visible_file_name=filename,
            caption=f"📊 Oylik hisobot ({today.month}-oy {today.year})"
        )
    except Exception as e:
        bot.send_message(message.chat.id, f"❌ Xatolik: {str(e)}")


@bot.message_handler(commands=['doctor'])
def doctor(message):
    user = message.from_user.id
    today = datetime.date.today()
    user_name = message.from_user.first_name
    
    for v in visits:
        if v["user"] == user and v["date"] == today:
            bot.send_message(message.chat.id, 
                           "⚠️ **Bugun allaqachon doktor kiritilgan!**\n\n"
                           "1 kunda faqat 1 marta doktor kiritish mumkin.",
                           parse_mode='Markdown')
            return

    visits.append({
        "user": user,
        "user_name": user_name,
        "date": today
    })

    today_count = sum(1 for v in visits if v["date"] == today and v["user"] == user)
    apteka_count = sum(1 for p in pharmacy_visits if p["date"] == today and p["user"] == user)

    success_text = f"""
✅ **Doktor tashrifi yozildi!**

👤 Agent: {user_name}
📅 Sana: {today}

🏥 Doktor: {today_count}/12
💊 Apteka: {apteka_count}/5

📊 KPI: {int((today_count/12)*100)}%
    """
    bot.send_message(message.chat.id, success_text, parse_mode='Markdown')


@bot.message_handler(commands=['apteka'])
def apteka(message):
    user = message.from_user.id
    today = datetime.date.today()
    user_name = message.from_user.first_name
    
    pharmacy_visits.append({
        "user": user,
        "user_name": user_name,
        "date": today
    })

    apteka_count = sum(1 for p in pharmacy_visits if p["date"] == today and p["user"] == user)
    doctor_count = sum(1 for v in visits if v["date"] == today and v["user"] == user)
    
    success_text = f"""
✅ **Apteka tashrifi yozildi!**

👤 Agent: {user_name}
📅 Sana: {today}

💊 Apteka: {apteka_count}/5
🏥 Doktor: {doctor_count}/12

📊 KPI: {int((doctor_count/12)*100)}%
    """
    bot.send_message(message.chat.id, success_text, parse_mode='Markdown')


@bot.message_handler(commands=['report'])
def report(message):
    today = datetime.date.today()
    user = message.from_user.id
    user_name = message.from_user.first_name
    
    doctor_count = sum(1 for v in visits if v["user"] == user and v["date"] == today)
    apteka_count = sum(1 for p in pharmacy_visits if p["user"] == user and p["date"] == today)
    
    doctor_kpi = min(100, (doctor_count / 12) * 100)
    apteka_kpi = min(100, (apteka_count / 5) * 100)
    total_kpi = int((doctor_kpi * 0.7) + (apteka_kpi * 0.3))
    
    doctor_progress = "🟩" * doctor_count + "⬜" * (12 - doctor_count)
    apteka_progress = "🟩" * apteka_count + "⬜" * (5 - apteka_count)
    
    report_text = f"""
📊 **KUNLIK HISOBOT**
👤 Agent: {user_name}
📅 Sana: {today}

🏥 **Doktorlar:** {doctor_count}/12
{doctor_progress}

💊 **Aptekalar:** {apteka_count}/5
{apteka_progress}

📈 **UMUMIY KPI: {total_kpi}%**

"""
    if doctor_count < 12 or apteka_count < 5:
        qolgan_doktor = 12 - doctor_count
        qolgan_apteka = 5 - apteka_count
        report_text += f"⏳ Qolgan: {qolgan_doktor} doktor, {qolgan_apteka} apteka"
    else:
        report_text += "🎉 Bugungi normani bajardingiz! Tabriklaymiz!"
    
    bot.send_message(message.chat.id, report_text, parse_mode='Markdown')


@bot.message_handler(commands=['plan'])
def plan(message):
    user_name = message.from_user.first_name
    plan_text = """
📝 **ERTANGI REJA TUZISH**

Ertangi kun uchun rejangizni yozing.
Format:
`Doktorlar soni | Aptekalar soni | Hudud`

Misol:
`12 | 5 | Chilonzor`
    """
    msg = bot.send_message(message.chat.id, plan_text, parse_mode='Markdown')
    bot.register_next_step_handler(msg, save_plan)


def save_plan(message):
    user_name = message.from_user.first_name
    plan_data = message.text
    
    tomorrow = datetime.date.today() + datetime.timedelta(days=1)
    
    success_text = f"""
✅ **ERTANGI REJA SAQLANDI!**

👤 Agent: {user_name}
📅 Sana: {tomorrow}

📋 Reja:
{plan_data}

Ertalab eslatma yuboriladi. Omad! 🚀
    """
    bot.send_message(message.chat.id, success_text, parse_mode='Markdown')
    
    try:
        manager_id = 6995658625
        bot.send_message(manager_id, 
                        f"📋 Yangi plan:\n👤 Agent: {user_name}\n📋 {plan_data}")
    except:
        pass


@bot.message_handler(commands=['stats'])
def stats(message):
    user = message.from_user.id
    user_name = message.from_user.first_name
    today = datetime.date.today()
    
    month_start = datetime.date(today.year, today.month, 1)
    month_visits = sum(1 for v in visits if v["user"] == user and v["date"] >= month_start)
    month_pharmacy = sum(1 for p in pharmacy_visits if p["user"] == user and p["date"] >= month_start)
    
    stats_text = f"""
📊 **OYLIK STATISTIKA**
👤 Agent: {user_name}
📅 Oy: {today.strftime('%B %Y')}

🏥 Jami doktor: {month_visits}
💊 Jami apteka: {month_pharmacy}
📈 O'rtacha doktor/kun: {month_visits / today.day:.1f}
📈 O'rtacha apteka/kun: {month_pharmacy / today.day:.1f}

🏆 Kategoriya: {
    "A (3+ marta)" if month_visits > 30 else 
    "B (2 marta)" if month_visits > 20 else 
    "C (1 marta)" if month_visits > 10 else 
    "D (0 marta)"
}
    """
    bot.send_message(message.chat.id, stats_text, parse_mode='Markdown')


@bot.message_handler(commands=['all_report'])
def all_report(message):
    if message.from_user.id != 6995658625:
        bot.send_message(message.chat.id, "⛔ Bu buyruq faqat manager uchun!")
        return
    
    today = datetime.date.today()
    agents_stats = {}
    
    for v in visits:
        if v["date"] == today:
            user_id = v["user"]
            if user_id not in agents_stats:
                agents_stats[user_id] = {
                    "name": v.get("user_name", f"User {user_id}"),
                    "doctor": 0,
                    "apteka": 0
                }
            agents_stats[user_id]["doctor"] += 1
    
    for p in pharmacy_visits:
        if p["date"] == today:
            user_id = p["user"]
            if user_id in agents_stats:
                agents_stats[user_id]["apteka"] += 1
    
    if not agents_stats:
        bot.send_message(message.chat.id, "📊 Bugun hech qanday tashrif yo'q")
        return
    
    report = "📊 **BUGUNGI UMUMIY HISOBOT**\n\n"
    for agent_id, data in agents_stats.items():
        doctor_kpi = min(100, (data["doctor"] / 12) * 100)
        apteka_kpi = min(100, (data["apteka"] / 5) * 100)
        total_kpi = int((doctor_kpi * 0.7) + (apteka_kpi * 0.3))
        
        report += f"👤 {data['name']}\n"
        report += f"   🏥 Doktor: {data['doctor']}/12\n"
        report += f"   💊 Apteka: {data['apteka']}/5\n"
        report += f"   📊 KPI: {total_kpi}%\n\n"
    
    bot.send_message(message.chat.id, report, parse_mode='Markdown')


@bot.message_handler(func=lambda message: message.text == '/')
def slash_command(message):
    commands_text = """
📋 **MAVJUD BUYRUQLAR:**

/start - Botni ishga tushirish
/help - Yordam va buyruqlar
/doctor - Doktor tashrifi qo'shish
/apteka - Apteka tashrifi qo'shish
/plan - Ertangi reja tuzish
/report - Bugungi hisobot
/stats - Oylik statistika

📊 **HISOBOTLAR:**
/monthly_excel - Excel formatda oylik hisobot
/monthly_word - Word formatda oylik hisobot
/all_report - Umumiy hisobot (manager)

ℹ️ Buyruqlarni tanlang va bosing.
    """
    bot.send_message(message.chat.id, commands_text, parse_mode='Markdown')


# Asosiy qism
if __name__ == "__main__":
    # Webhook sozlash uchun alohida thread
    threading.Thread(target=setup_webhook, daemon=True).start()
    
    # Flask serverni ishga tushirish
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port)
